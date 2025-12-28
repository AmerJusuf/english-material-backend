"""
Service for exporting materials to DOCX and PDF formats with professional styling.
"""
import io
import re
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, KeepTogether, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors


def add_border_to_paragraph(paragraph, color="4472C4", width=2):
    """Add a colored border around a paragraph."""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(width * 4))  # width in eighths of a point
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), color)
        pBdr.append(border)
    
    pPr.append(pBdr)


def add_shading_to_paragraph(paragraph, color="E7E6E6"):
    """Add background shading to a paragraph."""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)


class DocumentExporter:
    """Handles export of generated materials to various formats."""
    
    def __init__(self):
        pass
    
    def _parse_content(self, content: str) -> List[Dict]:
        """Parse content into structured sections with styling hints."""
        sections = []
        lines = content.split('\n')
        current_section = None
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Detect section types
            line_upper = line.upper()
            
            # Learning Objectives
            if 'LEARNING OBJECTIVE' in line_upper:
                if current_section:
                    sections.append(current_section)
                current_section = {'type': 'objectives', 'title': line, 'content': []}
            
            # Exercises
            elif re.match(r'EXERCISE\s+\d+', line_upper) or line_upper.startswith('EXERCISE:'):
                if current_section:
                    sections.append(current_section)
                current_section = {'type': 'exercise', 'title': line, 'content': []}
            
            # Major sections - Textbook structure
            elif any(keyword in line_upper for keyword in [
                'INTRODUCTION', 'WARM-UP', 'WARM UP', 'VOCABULARY', 'LANGUAGE FOCUS', 'GRAMMAR', 
                'READING TEXT', 'READING:', 'WRITING TASK', 'SPEAKING ACTIVITY', 'REVIEW', 
                'COMMUNICATION PRACTICE', 'SUMMARY', 'REFLECTION', 'REFLECTION SECTION',
                'GROUP WORK', 'DISCUSSION', 'DISCUSSION ACTIVITIES', 'ROLE-PLAY', 'ROLE PLAY',
                'DIALOGUES', 'DIALOGUE', 'INTERESTING FACTS', 'GAMES', 'EXERCISES'
            ]):
                if current_section:
                    sections.append(current_section)
                # Determine section type for better styling
                if 'INTRODUCTION' in line_upper:
                    section_type = 'introduction'
                elif 'SUMMARY' in line_upper:
                    section_type = 'summary'
                elif 'REFLECTION' in line_upper:
                    section_type = 'reflection'
                elif 'DIALOGUE' in line_upper:
                    section_type = 'dialogue'
                elif 'ROLE' in line_upper:
                    section_type = 'roleplay'
                elif 'GROUP WORK' in line_upper or 'DISCUSSION' in line_upper:
                    section_type = 'group_activity'
                else:
                    section_type = 'section'
                current_section = {'type': section_type, 'title': line, 'content': []}
            
            # Regular content
            else:
                if current_section is None:
                    current_section = {'type': 'paragraph', 'content': []}
                current_section['content'].append(line)
            
            i += 1
        
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def export_to_docx(self, material: Dict[str, Any]) -> io.BytesIO:
        """Export material to beautifully styled DOCX format."""
        doc = Document()
        
        # Set up default styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
        # Title with styling
        title = doc.add_heading(material['title'], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        title_run.font.bold = True
        
        doc.add_paragraph()
        
        # Process each chapter
        for chapter in material.get('chapters', []):
            # Chapter heading with blue background
            chapter_heading = doc.add_heading(
                f"Chapter {chapter['number']}: {chapter['title']}", 
                level=1
            )
            chapter_heading.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            chapter_heading.runs[0].font.size = Pt(18)
            add_shading_to_paragraph(chapter_heading, "4472C4")  # Blue background
            
            doc.add_paragraph()
            
            # Parse chapter content
            sections = self._parse_content(chapter['content'])
            
            for section in sections:
                section_type = section.get('type', 'paragraph')
                title = section.get('title', '')
                content = section.get('content', [])
                
                if section_type == 'objectives':
                    # Learning Objectives - in a light blue box
                    p = doc.add_paragraph()
                    run = p.add_run('📚 ' + title)
                    run.font.size = Pt(13)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 51, 102)
                    add_shading_to_paragraph(p, "D9E2F3")  # Light blue
                    add_border_to_paragraph(p, "4472C4", 2)
                    
                    for line in content:
                        if line:
                            p = doc.add_paragraph(line, style='List Bullet')
                            p.paragraph_format.left_indent = Inches(0.3)
                    doc.add_paragraph()
                
                elif section_type == 'exercise':
                    # Exercise - in a colored box
                    p = doc.add_paragraph()
                    run = p.add_run('✏️ ' + title)
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 102, 0)
                    add_shading_to_paragraph(p, "E2EFD9")  # Light green
                    add_border_to_paragraph(p, "70AD47", 2)
                    
                    for line in content:
                        if line:
                            # Check if it's an instruction line
                            if any(word in line.lower() for word in ['match', 'fill', 'complete', 'write', 'answer', 'choose']):
                                p = doc.add_paragraph(line)
                                run = p.runs[0]
                                run.font.italic = True
                                run.font.color.rgb = RGBColor(89, 89, 89)
                            else:
                                p = doc.add_paragraph(line)
                            
                            # Add extra space for answer lines
                            if '____' in line or 'answer:' in line.lower():
                                doc.add_paragraph('_' * 60)
                    
                    doc.add_paragraph()
                
                elif section_type == 'introduction':
                    # Introduction - prominent section
                    p = doc.add_paragraph()
                    run = p.add_run('📘 ' + title)
                    run.font.size = Pt(14)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 51, 102)
                    add_shading_to_paragraph(p, "D9E2F3")  # Light blue
                    add_border_to_paragraph(p, "4472C4", 2)
                    
                    for line in content:
                        if line:
                            p = doc.add_paragraph(line)
                            p.paragraph_format.line_spacing = 1.15
                    doc.add_paragraph()
                
                elif section_type == 'summary':
                    # Summary - distinct styling
                    p = doc.add_paragraph()
                    run = p.add_run('📋 ' + title)
                    run.font.size = Pt(13)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(102, 0, 102)
                    add_shading_to_paragraph(p, "E1D5E7")  # Light purple
                    add_border_to_paragraph(p, "7030A0", 2)
                    
                    for line in content:
                        if line:
                            p = doc.add_paragraph(line)
                            p.paragraph_format.line_spacing = 1.15
                    doc.add_paragraph()
                
                elif section_type == 'reflection':
                    # Reflection section - thoughtful styling
                    p = doc.add_paragraph()
                    run = p.add_run('💭 ' + title)
                    run.font.size = Pt(13)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 102, 51)
                    add_shading_to_paragraph(p, "D5E8D4")  # Light green
                    add_border_to_paragraph(p, "70AD47", 2)
                    
                    for line in content:
                        if line:
                            p = doc.add_paragraph(line)
                            p.paragraph_format.line_spacing = 1.15
                    doc.add_paragraph()
                
                elif section_type == 'dialogue':
                    # Dialogue - conversational styling
                    p = doc.add_paragraph()
                    run = p.add_run('💬 ' + title)
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(102, 51, 0)
                    add_shading_to_paragraph(p, "FFF2CC")  # Light yellow
                    add_border_to_paragraph(p, "FFC000", 2)
                    
                    for line in content:
                        if line:
                            p = doc.add_paragraph(line)
                            p.paragraph_format.left_indent = Inches(0.3)
                            p.paragraph_format.line_spacing = 1.2
                    doc.add_paragraph()
                
                elif section_type == 'roleplay':
                    # Role-play activity
                    p = doc.add_paragraph()
                    run = p.add_run('🎭 ' + title)
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(153, 0, 0)
                    add_shading_to_paragraph(p, "F4CCCC")  # Light red
                    add_border_to_paragraph(p, "C00000", 2)
                    
                    for line in content:
                        if line:
                            p = doc.add_paragraph(line)
                            p.paragraph_format.line_spacing = 1.15
                    doc.add_paragraph()
                
                elif section_type == 'group_activity':
                    # Group work or discussion
                    p = doc.add_paragraph()
                    run = p.add_run('👥 ' + title)
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 76, 153)
                    add_shading_to_paragraph(p, "D0E0F0")  # Light blue-gray
                    add_border_to_paragraph(p, "4F81BD", 2)
                    
                    for line in content:
                        if line:
                            p = doc.add_paragraph(line)
                            p.paragraph_format.line_spacing = 1.15
                    doc.add_paragraph()
                
                elif section_type == 'section':
                    # Major section heading - with colored background
                    p = doc.add_paragraph()
                    run = p.add_run('📖 ' + title)
                    run.font.size = Pt(13)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(102, 51, 0)
                    add_shading_to_paragraph(p, "FCE4D6")  # Light orange
                    add_border_to_paragraph(p, "C65911", 2)
                    
                    for line in content:
                        if line:
                            p = doc.add_paragraph(line)
                            p.paragraph_format.line_spacing = 1.15
                    doc.add_paragraph()
                
                else:
                    # Regular paragraph
                    for line in content:
                        if line:
                            p = doc.add_paragraph(line)
                            p.paragraph_format.line_spacing = 1.15
            
            # Add page break after each chapter except the last
            if chapter != material['chapters'][-1]:
                doc.add_page_break()
        
        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def export_to_pdf(self, material: Dict[str, Any]) -> io.BytesIO:
        """Export material to beautifully styled PDF format."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm,
        )
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#003366'),
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        # Chapter style
        chapter_style = ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.white,
            spaceAfter=12,
            spaceBefore=0,
            fontName='Helvetica-Bold',
            backColor=colors.HexColor('#4472C4'),
            borderPadding=8,
        )
        
        # Section header style
        section_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=colors.HexColor('#663300'),
            spaceAfter=8,
            spaceBefore=12,
            fontName='Helvetica-Bold',
            backColor=colors.HexColor('#FCE4D6'),
            borderPadding=6,
            leftIndent=10,
            rightIndent=10,
        )
        
        # Exercise style
        exercise_style = ParagraphStyle(
            'ExerciseHeader',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor('#006600'),
            spaceAfter=8,
            spaceBefore=10,
            fontName='Helvetica-Bold',
            backColor=colors.HexColor('#E2EFD9'),
            borderPadding=6,
            leftIndent=10,
            rightIndent=10,
        )
        
        # Objectives style
        objectives_style = ParagraphStyle(
            'ObjectivesHeader',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor('#003366'),
            spaceAfter=8,
            spaceBefore=10,
            fontName='Helvetica-Bold',
            backColor=colors.HexColor('#D9E2F3'),
            borderPadding=6,
            leftIndent=10,
            rightIndent=10,
        )
        
        # Body style
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            spaceAfter=6,
            alignment=TA_JUSTIFY,
            fontName='Helvetica'
        )
        
        # Instruction style
        instruction_style = ParagraphStyle(
            'Instruction',
            parent=styles['BodyText'],
            fontSize=10,
            spaceAfter=4,
            textColor=colors.HexColor('#595959'),
            fontName='Helvetica-Oblique',
            leftIndent=15,
        )
        
        # Introduction style
        introduction_style = ParagraphStyle(
            'Introduction',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=colors.HexColor('#003366'),
            spaceAfter=8,
            spaceBefore=12,
            fontName='Helvetica-Bold',
            backColor=colors.HexColor('#D9E2F3'),
            borderPadding=6,
            leftIndent=10,
            rightIndent=10,
        )
        
        # Summary style
        summary_style = ParagraphStyle(
            'Summary',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=colors.HexColor('#660066'),
            spaceAfter=8,
            spaceBefore=12,
            fontName='Helvetica-Bold',
            backColor=colors.HexColor('#E1D5E7'),
            borderPadding=6,
            leftIndent=10,
            rightIndent=10,
        )
        
        # Reflection style
        reflection_style = ParagraphStyle(
            'Reflection',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=colors.HexColor('#006633'),
            spaceAfter=8,
            spaceBefore=12,
            fontName='Helvetica-Bold',
            backColor=colors.HexColor('#D5E8D4'),
            borderPadding=6,
            leftIndent=10,
            rightIndent=10,
        )
        
        # Dialogue style
        dialogue_style = ParagraphStyle(
            'Dialogue',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor('#663300'),
            spaceAfter=6,
            spaceBefore=10,
            fontName='Helvetica-Bold',
            backColor=colors.HexColor('#FFF2CC'),
            borderPadding=6,
            leftIndent=10,
            rightIndent=10,
        )
        
        # Role-play style
        roleplay_style = ParagraphStyle(
            'Roleplay',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor('#990000'),
            spaceAfter=6,
            spaceBefore=10,
            fontName='Helvetica-Bold',
            backColor=colors.HexColor('#F4CCCC'),
            borderPadding=6,
            leftIndent=10,
            rightIndent=10,
        )
        
        # Group activity style
        group_activity_style = ParagraphStyle(
            'GroupActivity',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor('#004C99'),
            spaceAfter=6,
            spaceBefore=10,
            fontName='Helvetica-Bold',
            backColor=colors.HexColor('#D0E0F0'),
            borderPadding=6,
            leftIndent=10,
            rightIndent=10,
        )
        
        # Build document content
        story = []
        
        # Title
        story.append(Paragraph(material['title'], title_style))
        story.append(Spacer(1, 0.3 * inch))
        
        # Process each chapter
        for chapter in material.get('chapters', []):
            # Chapter title
            chapter_title = f"&nbsp;&nbsp;Chapter {chapter['number']}: {chapter['title']}&nbsp;&nbsp;"
            story.append(Paragraph(chapter_title, chapter_style))
            story.append(Spacer(1, 0.15 * inch))
            
            # Parse chapter content
            sections = self._parse_content(chapter['content'])
            
            for section in sections:
                section_type = section.get('type', 'paragraph')
                title = section.get('title', '')
                content = section.get('content', [])
                
                if section_type == 'objectives':
                    story.append(Paragraph(f"&nbsp;&nbsp;📚 {title}&nbsp;&nbsp;", objectives_style))
                    for line in content:
                        if line:
                            story.append(Paragraph(f"• {line}", body_style))
                    story.append(Spacer(1, 0.15 * inch))
                
                elif section_type == 'exercise':
                    story.append(Paragraph(f"&nbsp;&nbsp;✏️ {title}&nbsp;&nbsp;", exercise_style))
                    for line in content:
                        if line:
                            if any(word in line.lower() for word in ['match', 'fill', 'complete', 'write', 'answer', 'choose']):
                                story.append(Paragraph(line, instruction_style))
                            else:
                                story.append(Paragraph(line, body_style))
                            
                            if '____' in line or 'answer:' in line.lower():
                                story.append(Spacer(1, 0.2 * inch))
                                story.append(Paragraph('_' * 70, body_style))
                    story.append(Spacer(1, 0.15 * inch))
                
                elif section_type == 'introduction':
                    story.append(Paragraph(f"&nbsp;&nbsp;📘 {title}&nbsp;&nbsp;", introduction_style))
                    for line in content:
                        if line:
                            story.append(Paragraph(line, body_style))
                    story.append(Spacer(1, 0.15 * inch))
                
                elif section_type == 'summary':
                    story.append(Paragraph(f"&nbsp;&nbsp;📋 {title}&nbsp;&nbsp;", summary_style))
                    for line in content:
                        if line:
                            story.append(Paragraph(line, body_style))
                    story.append(Spacer(1, 0.15 * inch))
                
                elif section_type == 'reflection':
                    story.append(Paragraph(f"&nbsp;&nbsp;💭 {title}&nbsp;&nbsp;", reflection_style))
                    for line in content:
                        if line:
                            story.append(Paragraph(line, body_style))
                    story.append(Spacer(1, 0.15 * inch))
                
                elif section_type == 'dialogue':
                    story.append(Paragraph(f"&nbsp;&nbsp;💬 {title}&nbsp;&nbsp;", dialogue_style))
                    for line in content:
                        if line:
                            p = Paragraph(line, body_style)
                            p.style.leftIndent = 15
                            story.append(p)
                    story.append(Spacer(1, 0.15 * inch))
                
                elif section_type == 'roleplay':
                    story.append(Paragraph(f"&nbsp;&nbsp;🎭 {title}&nbsp;&nbsp;", roleplay_style))
                    for line in content:
                        if line:
                            story.append(Paragraph(line, body_style))
                    story.append(Spacer(1, 0.15 * inch))
                
                elif section_type == 'group_activity':
                    story.append(Paragraph(f"&nbsp;&nbsp;👥 {title}&nbsp;&nbsp;", group_activity_style))
                    for line in content:
                        if line:
                            story.append(Paragraph(line, body_style))
                    story.append(Spacer(1, 0.15 * inch))
                
                elif section_type == 'section':
                    story.append(Paragraph(f"&nbsp;&nbsp;📖 {title}&nbsp;&nbsp;", section_style))
                    for line in content:
                        if line:
                            story.append(Paragraph(line, body_style))
                    story.append(Spacer(1, 0.15 * inch))
                
                else:
                    for line in content:
                        if line:
                            story.append(Paragraph(line, body_style))
            
            # Add page break after each chapter except the last
            if chapter != material['chapters'][-1]:
                story.append(PageBreak())
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer


# Singleton instance
document_exporter = DocumentExporter()
